#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Flask API сервис для генерации Word документов (заявлений)
Запуск: python flask_api_server.py
URL: http://localhost:5000
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import io
import os
from datetime import datetime
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)  # Разрешаем CORS для n8n


class ArizaDocumentGenerator:
    """Генератор документов заявлений"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Настройка стилей документа"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)
            section.left_margin = Inches(1.18)
            section.right_margin = Inches(0.59)
        
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
    
    def _add_right_aligned_paragraph(self, text, bold=False):
        """Добавить абзац с выравниванием по правому краю"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = bold
        return p
    
    def _add_center_paragraph(self, text, bold=True):
        """Добавить абзац с выравниванием по центру"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = bold
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(18)
        return p
    
    def _add_body_paragraph(self, text, first_line_indent=True):
        """Добавить основной текст с выравниванием по ширине"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        if first_line_indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        return p
    
    def _add_signature_line(self, left_text, right_text):
        """Добавить строку с подписью"""
        p = self.doc.add_paragraph()
        
        run_left = p.add_run(left_text)
        run_left.font.name = 'Times New Roman'
        run_left.font.size = Pt(14)
        
        # Добавляем пробелы для выравнивания
        spacing = ' ' * 40
        run_spacing = p.add_run(spacing)
        
        run_right = p.add_run(right_text)
        run_right.font.name = 'Times New Roman'
        run_right.font.size = Pt(14)
        
        return p
    
    def parse_ariza_text(self, text):
        """Парсинг текста заявления"""
        lines = text.split('\n')
        
        header_lines = []
        body_lines = []
        appendix_lines = []
        footer_date = ""
        footer_signature = ""
        
        current_section = 'header'
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                continue
            
            if 'А Р И З А' in stripped or 'АРИЗА' in stripped:
                current_section = 'body'
                continue
            
            if stripped.startswith('Илова:'):
                current_section = 'appendix'
            
            if re.search(r'\d{2}\.\d{2}\.\d{4}', stripped):
                footer_date = stripped
                current_section = 'footer'
                continue
            
            if current_section == 'footer' or 'Адвокат' in stripped or 'Имзо' in stripped:
                if not footer_date and stripped:
                    footer_date = stripped
                elif stripped:
                    footer_signature = stripped
                continue
            
            if current_section == 'header':
                header_lines.append(stripped)
            elif current_section == 'body':
                body_lines.append(stripped)
            elif current_section == 'appendix':
                appendix_lines.append(stripped)
        
        return {
            'header': header_lines,
            'body': body_lines,
            'appendix': appendix_lines,
            'date': footer_date or datetime.now().strftime('%d.%m.%Y йил'),
            'signature': footer_signature or '[Имзо]'
        }
    
    def generate_document(self, ariza_data):
        """Генерация полного документа"""
        # Шапка
        for line in ariza_data['header']:
            self._add_right_aligned_paragraph(line, bold=('судига' in line.lower()))
        
        self.doc.add_paragraph()
        
        # Заголовок
        self._add_center_paragraph('А Р И З А', bold=True)
        
        # Основной текст
        for i, paragraph in enumerate(ariza_data['body']):
            self._add_body_paragraph(paragraph, first_line_indent=(i == 0))
        
        # Приложения
        if ariza_data['appendix']:
            self.doc.add_paragraph()
            for line in ariza_data['appendix']:
                self._add_body_paragraph(line, first_line_indent=False)
        
        # Дата и подпись
        self.doc.add_paragraph()
        self._add_signature_line(ariza_data['date'], ariza_data['signature'])
    
    def save_to_bytes(self):
        """Сохранить документ в байты (для отправки по API)"""
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        return file_stream


@app.route('/health', methods=['GET'])
def health_check():
    """Проверка работоспособности сервиса"""
    return jsonify({
        'status': 'ok',
        'service': 'Ariza Document Generator',
        'version': '1.0.0'
    })


@app.route('/generate-ariza', methods=['POST'])
def generate_ariza():
    """
    Основной эндпоинт для генерации документа
    
    Request JSON:
    {
        "text": "Полный текст заявления",
        "filename": "ariza.docx" (опционально)
    }
    
    Response: Word документ (.docx)
    """
    try:
        data = request.get_json()
        
        if not data or 'text' not in data:
            return jsonify({'error': 'Поле "text" обязательно'}), 400
        
        document_text = data['text']
        filename = data.get('filename', f'ariza_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        
        logger.info(f'Получен запрос на генерацию документа: {filename}')
        
        # Создаём генератор и парсим текст
        generator = ArizaDocumentGenerator()
        parsed_data = generator.parse_ariza_text(document_text)
        
        logger.info(f'Текст успешно распарсен. Секций: header={len(parsed_data["header"])}, body={len(parsed_data["body"])}')
        
        # Генерируем документ
        generator.generate_document(parsed_data)
        
        # Сохраняем в байты
        file_stream = generator.save_to_bytes()
        
        logger.info(f'Документ успешно сгенерирован: {filename}')
        
        # Отправляем файл
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        logger.error(f'Ошибка генерации документа: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/generate-ariza-base64', methods=['POST'])
def generate_ariza_base64():
    """
    Альтернативный эндпоинт - возвращает документ в Base64
    Полезно для n8n, если есть проблемы с бинарными данными
    
    Response JSON:
    {
        "filename": "ariza.docx",
        "content": "base64_encoded_content",
        "mimeType": "application/vnd..."
    }
    """
    try:
        data = request.get_json()
        
        if not data or 'text' not in data:
            return jsonify({'error': 'Поле "text" обязательно'}), 400
        
        document_text = data['text']
        filename = data.get('filename', f'ariza_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        
        # Генерируем документ
        generator = ArizaDocumentGenerator()
        parsed_data = generator.parse_ariza_text(document_text)
        generator.generate_document(parsed_data)
        file_stream = generator.save_to_bytes()
        
        # Конвертируем в Base64
        import base64
        file_content = file_stream.read()
        base64_content = base64.b64encode(file_content).decode('utf-8')
        
        logger.info(f'Документ сгенерирован и закодирован в Base64: {filename}')
        
        return jsonify({
            'filename': filename,
            'content': base64_content,
            'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'size': len(file_content)
        })
    
    except Exception as e:
        logger.error(f'Ошибка генерации документа: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/parse-ariza', methods=['POST'])
def parse_ariza():
    """
    Утилита для тестирования парсинга текста
    
    Request JSON:
    {
        "text": "Текст заявления"
    }
    
    Response JSON:
    {
        "header": [...],
        "body": [...],
        "appendix": [...],
        "date": "...",
        "signature": "..."
    }
    """
    try:
        data = request.get_json()
        
        if not data or 'text' not in data:
            return jsonify({'error': 'Поле "text" обязательно'}), 400
        
        generator = ArizaDocumentGenerator()
        parsed_data = generator.parse_ariza_text(data['text'])
        
        return jsonify(parsed_data)
    
    except Exception as e:
        logger.error(f'Ошибка парсинга: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('DEBUG', 'False').lower() == 'true'
    
    logger.info(f'🚀 Запуск Flask API сервера на порту {port}')
    logger.info(f'📝 Эндпоинты:')
    logger.info(f'   - GET  /health')
    logger.info(f'   - POST /generate-ariza')
    logger.info(f'   - POST /generate-ariza-base64')
    logger.info(f'   - POST /parse-ariza')
    
    app.run(host='0.0.0.0', port=port, debug=debug)
