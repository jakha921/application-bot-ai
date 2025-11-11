#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для генерации Word документов (заявлений) 
с правильным форматированием по образцу
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from datetime import datetime


class ArizaDocumentGenerator:
    """Генератор документов заявлений для Узбекистана"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Настройка стилей документа"""
        # Поля документа
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)
            section.left_margin = Inches(1.18)
            section.right_margin = Inches(0.59)
        
        # Стиль по умолчанию
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
    
    def _add_right_aligned_paragraph(self, text, bold=False, italic=False):
        """Добавить абзац с выравниванием по правому краю"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = bold
        run.italic = italic
        return p
    
    def _add_center_paragraph(self, text, bold=True, spacing_before=12, spacing_after=12):
        """Добавить абзац с выравниванием по центру"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = bold
        p.paragraph_format.space_before = Pt(spacing_before)
        p.paragraph_format.space_after = Pt(spacing_after)
        return p
    
    def _add_body_paragraph(self, text, first_line_indent=True):
        """Добавить основной текст с выравниванием по ширине"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        if first_line_indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        return p
    
    def _add_signature_line(self, left_text, right_text):
        """Добавить строку с подписью (дата слева, подпись справа)"""
        p = self.doc.add_paragraph()
        
        # Левая часть (дата)
        run_left = p.add_run(left_text)
        run_left.font.name = 'Times New Roman'
        run_left.font.size = Pt(14)
        
        # Добавляем табуляцию для выравнивания
        run_left.add_tab()
        run_left.add_tab()
        run_left.add_tab()
        
        # Правая часть (подпись)
        run_right = p.add_run(right_text)
        run_right.font.name = 'Times New Roman'
        run_right.font.size = Pt(14)
        
        return p
    
    def parse_ariza_text(self, text):
        """
        Парсинг текста заявления и извлечение структурных элементов
        
        Структура:
        - Шапка (правый верхний угол)
        - Заголовок (А Р И З А)
        - Основной текст
        - Приложения (Илова)
        - Дата и подпись
        """
        lines = text.split('\n')
        
        # Инициализация секций
        header_lines = []
        title_found = False
        body_lines = []
        appendix_lines = []
        footer_date = ""
        footer_signature = ""
        
        current_section = 'header'
        
        for line in lines:
            stripped = line.strip()
            
            # Пропускаем пустые строки в начале
            if not stripped and current_section == 'header':
                continue
            
            # Определяем заголовок
            if 'А Р И З А' in stripped or 'АРИЗА' in stripped:
                title_found = True
                current_section = 'body'
                continue
            
            # Определяем секцию приложений
            if stripped.startswith('Илова:') or stripped.startswith('Аризага қуйидагилар'):
                current_section = 'appendix'
            
            # Определяем дату (формат: DD.MM.YYYY или просто дата)
            if re.search(r'\d{2}\.\d{2}\.\d{4}', stripped):
                footer_date = stripped
                current_section = 'footer'
                continue
            
            # Определяем подпись (обычно после даты или содержит "Адвокат", "Имзо")
            if current_section == 'footer' or 'Адвокат' in stripped or 'Имзо' in stripped:
                if stripped and not footer_date:
                    footer_date = stripped
                elif stripped:
                    footer_signature = stripped
                continue
            
            # Распределяем строки по секциям
            if current_section == 'header':
                if stripped:
                    header_lines.append(stripped)
            elif current_section == 'body':
                if stripped:
                    body_lines.append(stripped)
            elif current_section == 'appendix':
                if stripped:
                    appendix_lines.append(stripped)
        
        return {
            'header': header_lines,
            'body': body_lines,
            'appendix': appendix_lines,
            'date': footer_date or datetime.now().strftime('%d.%m.%Y йил'),
            'signature': footer_signature or '[Имзо]'
        }
    
    def generate_document(self, ariza_data):
        """
        Генерация полного документа заявления
        
        Args:
            ariza_data: dict с ключами header, body, appendix, date, signature
        """
        # 1. ШАПКА (правый верхний угол)
        for line in ariza_data['header']:
            self._add_right_aligned_paragraph(line, bold=('судига' in line.lower()))
        
        # Пустая строка после шапки
        self.doc.add_paragraph()
        
        # 2. ЗАГОЛОВОК
        self._add_center_paragraph('А Р И З А', bold=True, spacing_before=18, spacing_after=18)
        
        # 3. ОСНОВНОЙ ТЕКСТ
        for i, paragraph in enumerate(ariza_data['body']):
            # Первый параграф обычно с отступом
            self._add_body_paragraph(paragraph, first_line_indent=(i == 0))
        
        # Пустая строка перед приложениями
        if ariza_data['appendix']:
            self.doc.add_paragraph()
        
        # 4. ПРИЛОЖЕНИЯ
        if ariza_data['appendix']:
            for line in ariza_data['appendix']:
                self._add_body_paragraph(line, first_line_indent=False)
        
        # Пустая строка перед датой и подписью
        self.doc.add_paragraph()
        
        # 5. ДАТА И ПОДПИСЬ
        self._add_signature_line(ariza_data['date'], ariza_data['signature'])
    
    def save(self, filename):
        """Сохранить документ"""
        self.doc.save(filename)
        print(f"✅ Документ сохранён: {filename}")


def generate_ariza_from_text(text, output_filename='ariza.docx'):
    """
    Главная функция для генерации документа из текста
    
    Args:
        text: Полный текст заявления
        output_filename: Имя выходного файла
    
    Returns:
        str: Путь к созданному файлу
    """
    generator = ArizaDocumentGenerator()
    parsed_data = generator.parse_ariza_text(text)
    generator.generate_document(parsed_data)
    generator.save(output_filename)
    return output_filename


# ============ ПРИМЕР ИСПОЛЬЗОВАНИЯ ============

if __name__ == "__main__":
    # Пример текста заявления (из вашего образца)
    sample_text = """
Жиноят ишлари бўйича Навоий шаҳар судига
Жиноят кодекси (кейинги ўринларда ЖК деб
аталади)нинг 167-модда учинчи қисми ва бошқа
моддалари билан айбланган "Иброҳим экспорт
кластер" масъулияти чекланган жамияти раҳбари
Холмуродов Иброҳим Шомирза ўғлининг
ҳимоячиси нисбатан жиноят иши бўйича "Баҳс-
Навоий" адвокатлар ҳайъати адвокати
Д.Р.Ражабов томонидан
Манзил: Навоий шаҳри, "Ўзбекистон" кўчаси, 14-уй
адвокат тел: (90) 501-52-47

А Р И З А

Сиздан ЖКнинг 257-модда иккинчи қисми "а" банди билан жиноят ишлари
бўйича Навоий шаҳар судининг 29.04.2025 йилдаги ҳукми билан судланган
Шомуродов Фарҳод Барноевичга нисбатан 1-2101-2501/108-сонли жиноят иши
бўйича тўпланган иш ҳужжатлари билан танишиш ва улардан фотонусхалар
олишга рухсат беришингизни сўрайман.

Илова: 27-сонли 14.10.2025 йилдаги адвокат томонидан ишни олиб боришга
ордер 1 варақдан иборат.

Адвокат Д.Ражабов
14.10.2025 йил
    """
    
    # Генерируем документ
    output_file = generate_ariza_from_text(sample_text, 'sample_ariza.docx')
    print(f"\n🎉 Готово! Проверьте файл: {output_file}")
