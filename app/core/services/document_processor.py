"""
Document Processing Service
Handles text extraction from PDF, DOCX files and content processing
"""
import logging
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Custom exception for document processing errors"""
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using pdfplumber (more reliable than PyPDF2)
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
        
    Raises:
        DocumentProcessingError: If extraction fails
    """
    try:
        import pdfplumber
        
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        if not text_parts:
            raise DocumentProcessingError("No text could be extracted from PDF")
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
        return full_text
        
    except ImportError:
        # Fallback to PyPDF2 if pdfplumber not available
        logger.warning("pdfplumber not available, falling back to PyPDF2")
        return _extract_text_from_pdf_pypdf2(file_path)
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise DocumentProcessingError(f"Failed to extract text from PDF: {str(e)}")


def _extract_text_from_pdf_pypdf2(file_path: str) -> str:
    """Fallback PDF extraction using PyPDF2"""
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        if not text_parts:
            raise DocumentProcessingError("No text could be extracted from PDF")
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise DocumentProcessingError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
        
    Raises:
        DocumentProcessingError: If extraction fails
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    table_texts.append(row_text)
        
        # Combine all text
        all_text = "\n\n".join(paragraphs)
        if table_texts:
            all_text += "\n\nTables:\n" + "\n".join(table_texts)
        
        if not all_text.strip():
            raise DocumentProcessingError("No text could be extracted from DOCX")
        
        logger.info(f"Successfully extracted {len(all_text)} characters from DOCX")
        return all_text
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise DocumentProcessingError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """
    Extract text from file based on type
    
    Args:
        file_path: Path to file
        file_type: Type of file ('pdf', 'docx', 'text', 'url')
        
    Returns:
        Extracted text content
        
    Raises:
        DocumentProcessingError: If extraction fails or unsupported type
    """
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    elif file_type == 'text':
        # For text files, just read the content
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise DocumentProcessingError(f"Failed to read text file: {str(e)}")
    else:
        raise DocumentProcessingError(f"Unsupported file type: {file_type}")


def process_knowledge_file(knowledge_file) -> Tuple[bool, Optional[str]]:
    """
    Process a KnowledgeBaseFile instance - extract text and update status
    
    Args:
        knowledge_file: KnowledgeBaseFile model instance
        
    Returns:
        Tuple of (success: bool, error_message: Optional[str])
    """
    from django.utils import timezone
    
    try:
        # If it's a text type, content should already be set
        if knowledge_file.file_type == 'text':
            if not knowledge_file.content:
                return False, "Text content is required for text type"
            knowledge_file.status = 'ready'
            knowledge_file.processed_at = timezone.now()
            knowledge_file.save()
            
            # Generate embeddings for text
            try:
                from core.services.embeddings_service import (
                    process_knowledge_file_embeddings
                )
                success, error = process_knowledge_file_embeddings(
                    knowledge_file
                )
                if not success:
                    logger.warning(
                        f"Embeddings generation failed: {error}"
                    )
            except Exception as emb_error:
                logger.warning(
                    f"Embeddings generation skipped: {str(emb_error)}"
                )
            
            return True, None
        
        # For URL type - would need URL fetching implementation
        if knowledge_file.file_type == 'url':
            return False, "URL processing not yet implemented"
        
        # For PDF/DOCX - extract text from file
        if not knowledge_file.file:
            return False, f"File is required for {knowledge_file.file_type} type"
        
        file_path = knowledge_file.file.path
        
        # Extract text
        logger.info(f"Processing {knowledge_file.file_type} file: {file_path}")
        extracted_text = extract_text_from_file(file_path, knowledge_file.file_type)
        
        # Update model
        knowledge_file.content = extracted_text
        knowledge_file.file_size = Path(file_path).stat().st_size
        knowledge_file.status = 'ready'
        knowledge_file.processed_at = timezone.now()
        knowledge_file.processing_error = None
        knowledge_file.save()
        
        # Generate embeddings for RAG (optional, can fail without blocking)
        try:
            from core.services.embeddings_service import (
                process_knowledge_file_embeddings
            )
            success, error = process_knowledge_file_embeddings(knowledge_file)
            if not success:
                logger.warning(
                    f"Embeddings generation failed for file "
                    f"{knowledge_file.id}: {error}"
                )
        except Exception as emb_error:
            logger.warning(
                f"Embeddings generation skipped for file "
                f"{knowledge_file.id}: {str(emb_error)}"
            )
        
        logger.info(f"Successfully processed knowledge file {knowledge_file.id}")
        return True, None
        
    except DocumentProcessingError as e:
        error_msg = str(e)
        logger.error(
            f"Processing failed for knowledge file "
            f"{knowledge_file.id}: {error_msg}"
        )
        knowledge_file.status = 'error'
        knowledge_file.processing_error = error_msg
        knowledge_file.save()
        return False, error_msg
        
    except Exception as e:
        error_msg = f"Unexpected error: {str(e)}"
        logger.exception(
            f"Unexpected error processing knowledge file "
            f"{knowledge_file.id}"
        )
        knowledge_file.status = 'error'
        knowledge_file.processing_error = error_msg
        knowledge_file.save()
        return False, error_msg
