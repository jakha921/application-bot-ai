"""
Word Document Generator for Ariza
Based on docs/Flask API Server for Telegram Workflow.py
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import io
from datetime import datetime
import logging

logger = logging.getLogger(__name__)


class ArizaDocumentGenerator:
    """Generator for Uzbek legal documents (Ariza)"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Configure document styles"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)
            section.left_margin = Inches(1.18)
            section.right_margin = Inches(0.59)
        
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
    
    def _add_right_aligned_paragraph(self, text, bold=False):
        """Add right-aligned paragraph"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = bold
        return p
    
    def _add_center_paragraph(self, text, bold=True):
        """Add center-aligned paragraph"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = bold
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(18)
        return p
    
    def _add_body_paragraph(self, text, first_line_indent=True):
        """Add justified body paragraph"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        if first_line_indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        return p
    
    def _add_signature_line(self, left_text, right_text):
        """Add signature line with date and signature"""
        p = self.doc.add_paragraph()
        
        run_left = p.add_run(left_text)
        run_left.font.name = 'Times New Roman'
        run_left.font.size = Pt(14)
        
        # Add spacing
        spacing = ' ' * 40
        p.add_run(spacing)
        
        run_right = p.add_run(right_text)
        run_right.font.name = 'Times New Roman'
        run_right.font.size = Pt(14)
        
        return p
    
    def parse_ariza_text(self, text):
        """
        Parse ariza text and extract structural elements
        
        Args:
            text: Full ariza text
        
        Returns:
            dict: Parsed structure with header, body, appendix, date, signature
        """
        lines = text.split('\n')
        
        header_lines = []
        body_lines = []
        appendix_lines = []
        footer_date = ""
        footer_signature = ""
        
        current_section = 'header'
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                continue
            
            # Detect title
            if 'А Р И З А' in stripped or 'АРИЗА' in stripped:
                current_section = 'body'
                continue
            
            # Detect appendix
            if stripped.startswith('Илова:'):
                current_section = 'appendix'
            
            # Detect date
            if re.search(r'\d{2}\.\d{2}\.\d{4}', stripped):
                footer_date = stripped
                current_section = 'footer'
                continue
            
            # Detect signature
            if current_section == 'footer' or 'Адвокат' in stripped or 'Имзо' in stripped:
                if not footer_date and stripped:
                    footer_date = stripped
                elif stripped:
                    footer_signature = stripped
                continue
            
            # Distribute lines to sections
            if current_section == 'header':
                header_lines.append(stripped)
            elif current_section == 'body':
                body_lines.append(stripped)
            elif current_section == 'appendix':
                appendix_lines.append(stripped)
        
        return {
            'header': header_lines,
            'body': body_lines,
            'appendix': appendix_lines,
            'date': footer_date or datetime.now().strftime('%d.%m.%Y йил'),
            'signature': footer_signature or '[Имзо]'
        }
    
    def generate_document(self, ariza_data):
        """
        Generate complete document
        
        Args:
            ariza_data: dict from parse_ariza_text()
        """
        # Header
        for line in ariza_data['header']:
            self._add_right_aligned_paragraph(
                line,
                bold=('судига' in line.lower())
            )
        
        self.doc.add_paragraph()
        
        # Title
        self._add_center_paragraph('А Р И З А', bold=True)
        
        # Body
        for i, paragraph in enumerate(ariza_data['body']):
            self._add_body_paragraph(paragraph, first_line_indent=(i == 0))
        
        # Appendix
        if ariza_data['appendix']:
            self.doc.add_paragraph()
            for line in ariza_data['appendix']:
                self._add_body_paragraph(line, first_line_indent=False)
        
        # Date and signature
        self.doc.add_paragraph()
        self._add_signature_line(ariza_data['date'], ariza_data['signature'])
    
    def save_to_bytes(self):
        """
        Save document to bytes for file storage
        
        Returns:
            io.BytesIO: Document as bytes
        """
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    
    def save_to_file(self, filename):
        """Save document to file"""
        self.doc.save(filename)
        logger.info(f"Document saved: {filename}")


def generate_ariza_document(document_text):
    """
    Main function to generate ariza document
    
    Args:
        document_text: Full text of ariza
    
    Returns:
        io.BytesIO: Generated Word document as bytes
    """
    try:
        generator = ArizaDocumentGenerator()
        parsed_data = generator.parse_ariza_text(document_text)
        generator.generate_document(parsed_data)
        
        logger.info("Ariza document generated successfully")
        return generator.save_to_bytes()
        
    except Exception as e:
        logger.error(f"Error generating ariza document: {e}")
        raise
